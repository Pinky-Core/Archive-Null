from copy import deepcopy
from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "Assets" / "ArchiveNull" / "Documentation"
TEMPLATE = next(DOCS.glob("Blanco*.docx"))
OUTPUT = DOCS / "GDD_Archive_NULL-V4.docx"
CONCEPT_IMAGE = ROOT / "Assets" / "ArchiveNull" / "Art" / "Brand" / "archive-null-abstract-reference.png"
SITUATION_SKETCHES = ROOT / "Assets" / "ArchiveNull" / "Art" / "Brand" / "archive-null-situation-sketches.png"


def normalize(value: str) -> str:
    return re.sub(r"\s+", " ", value).strip()


def find_paragraph(document: Document, prefix: str):
    for paragraph in document.paragraphs:
        if normalize(paragraph.text).startswith(prefix):
            return paragraph
    raise ValueError(f"Paragraph not found: {prefix}")


def clear_paragraph(paragraph):
    paragraph._element.clear_content()


def set_paragraph_text(paragraph, text: str, bold: bool = False, size: float = 10.5):
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)


def remove_section_paragraphs(document: Document, heading_prefix: str, next_heading_prefix: str):
    start = find_paragraph(document, heading_prefix)
    end = find_paragraph(document, next_heading_prefix)
    node = start._element.getnext()
    while node is not None and node is not end._element:
        following = node.getnext()
        if node.tag.endswith("}p"):
            node.getparent().remove(node)
        node = following
    return start, end


def insert_paragraph_before(reference, text: str, bold: bool = False, size: float = 10.5, bullet: bool = False):
    paragraph = reference._parent.add_paragraph()
    paragraph._element.getparent().remove(paragraph._element)
    reference._element.addprevious(paragraph._element)
    if bullet:
        text = f"• {text}"
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.05
    return paragraph


def replace_section(document: Document, heading: str, next_heading: str, blocks):
    _, end = remove_section_paragraphs(document, heading, next_heading)
    for block in blocks:
        if isinstance(block, tuple):
            kind, text = block
            insert_paragraph_before(end, text, bold=kind == "bold", bullet=kind == "bullet")
        else:
            insert_paragraph_before(end, block)


def set_cell(cell, text: str, bold: bool = False, size: float = 9.0):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    cell.vertical_alignment = 1


def fill_cover(document: Document):
    values = {
        "Nombre del juego:": "Nombre del juego: Archive: NULL",
        "Alumno:": "Alumno: Luca Caporaloni",
        "Materia:": "Materia: Proyecto de Integración: Desarrollo e Implementación de Videojuegos",
        "Fecha:": "Fecha: 12/07/2026",
    }
    for prefix, value in values.items():
        set_paragraph_text(find_paragraph(document, prefix), value, size=11)

    revision = document.tables[0]
    set_cell(revision.rows[1].cells[0], "12/07/2026", size=8.5)
    set_cell(
        revision.rows[1].cells[1],
        "Reescritura integral desde la plantilla original. Se completaron ideación, framework creativo, mecánicas, niveles, UX y código sin anexos ni contenido duplicado.",
        size=8.5,
    )


def fill_situations(document: Document):
    situations = [
        (
            "Revelado UV localizado",
            "El jugador barre el azucarero con la luz UV. Solo se revela la porción alcanzada por el cono y la intensidad acumulada; mover el haz demasiado rápido deja zonas incompletas. Una máscara de revelado conserva el recorrido y permite fotografiar el patrón final. La situación combina orientación del spotlight, tiempo de exposición y escritura progresiva sobre una textura.",
        ),
        (
            "Dos aplicaciones, una contradicción",
            "En el teléfono, el jugador compara el mensaje final con conversaciones anteriores y el registro de llamadas. El vocabulario, la hora y una llamada borrada contradicen la supuesta despedida. Cada app registra evidencia independiente y una regla de cruce habilita una conclusión nueva solo cuando ambas fueron consultadas.",
        ),
        (
            "Reconstrucción de la puerta cerrada",
            "El jugador encuentra hilo de obra, marcas en el picaporte y un plano de reforma. En una vista de reconstrucción debe ordenar puntos de paso para reproducir cómo se cerró la puerta desde afuera. El sistema valida longitud, orden espacial y colisiones del recorrido, no una respuesta de opción múltiple.",
        ),
        (
            "La memoria cambia de significado",
            "Una marca junto a la ventana no es registrable durante la primera visita. Después de hallar una herramienta compatible en la oficina del contratista, al volver a la casa aparece un nuevo comentario y la marca puede fotografiarse. El objeto conserva estado, pero su interpretación depende del conocimiento persistente del jugador.",
        ),
        (
            "Conexión incorrecta con respuesta narrativa",
            "En la pizarra se pueden unir libremente dos evidencias. Si la relación no demuestra método, acceso, motivo o tiempo, la conexión permanece como hipótesis y el detective explica qué dato falta. Un evaluador semántico compara categorías y relaciones permitidas sin bloquear la experimentación.",
        ),
        (
            "Comparación forense de documentos",
            "Dos facturas parecen iguales. El jugador las superpone sobre una mesa de luz y desplaza una capa hasta alinear sellos y firmas. Debe marcar tres diferencias: numeración repetida, fecha alterada y monto incompatible. Las zonas detectables se evalúan en coordenadas locales y alimentan una única evidencia compuesta.",
        ),
        (
            "El estado físico aporta una hora",
            "Una taza conserva calor y un trapo está húmedo. El jugador puede inspeccionarlos al entrar o después de explorar otros cuartos; sus valores cambian con el tiempo de la memoria. La combinación no identifica al culpable, pero acota el orden de preparación, limpieza y muerte.",
        ),
        (
            "Fotografía asistida por encuadre",
            "La cámara no depende de acertar un único píxel. El sistema evalúa distancia, visibilidad y cuánto ocupa la evidencia dentro del encuadre; la retícula cambia cuando la toma es válida. Una foto parcial puede registrarse con menor calidad y pedir una segunda toma que muestre el contexto completo.",
        ),
        (
            "Una evidencia incrimina y otra limita",
            "Una firma falsificada señala a Nicolás, pero un comprobante temporal lo ubica lejos durante la muerte. La pizarra admite ambas afirmaciones sin cancelar ninguna: delito documental y autoría del homicidio son conclusiones separadas. El jugador debe sostener hipótesis simultáneas hasta resolver sus requisitos.",
        ),
        (
            "Acusación construida, no adivinada",
            "El final exige seleccionar responsable y presentar una cadena de método, acceso, motivo, tiempo y manipulación de escena. El sistema recorre el grafo de evidencias; si falta un eslabón, devuelve una objeción específica y permite continuar investigando. Elegir el nombre correcto sin argumento no completa el caso.",
        ),
    ]

    table = document.tables[1]
    for index, (title, description) in enumerate(situations):
        row = table.rows[index]
        set_cell(row.cells[0], f"Situación {index + 1}\n{title}", bold=True, size=8.5)
        set_cell(row.cells[1], description, size=8.2)


def fill_framework(document: Document):
    table = document.tables[2]
    if len(table.rows) == 6:
        nuclear_values_row = table.rows[5]._tr
        abstraction_row = table.add_row()._tr
        nuclear_values_row.addprevious(abstraction_row)
    set_cell(
        table.rows[0].cells[0],
        "FANTASÍA: ser un detective de reconstrucción forense que entra en memorias incompletas, distingue hechos de interpretaciones y presenta una acusación demostrable.",
        bold=True,
        size=9.2,
    )
    set_cell(
        table.rows[1].cells[0],
        "El jugador no recibe la verdad terminada: observa, registra, contrasta y decide qué explicación resiste todas las contradicciones.",
        size=8.8,
    )

    headers = [
        "ABSTRACCIONES CLAVE",
        "PILAR 1\nObservación forense",
        "PILAR 2\nEvidencia con contexto",
        "PILAR 3\nDeducción argumentada",
    ]
    for column, value in enumerate(headers):
        set_cell(table.rows[2].cells[column], value, bold=True, size=8.3)

    rows = [
        (
            "Memoria investigable\nEscena reconstruida que puede revisitarse y cambiar de significado.",
            "Explorar espacios, inspeccionar estados y detectar anomalías visuales, temporales o sonoras.",
            "Fotografiar, revelar o consultar una pista conserva nombre, descripción, procedencia y estado.",
            "Relacionar método, acceso, motivo y tiempo hasta formar una explicación verificable.",
        ),
        (
            "Expediente y evidencia\nEl expediente aporta contexto inicial; la evidencia demuestra o contradice hechos.",
            "La primera explicación siempre puede ser incompleta o deliberadamente construida.",
            "Libreta y galería cumplen funciones distintas: reflexión libre frente a registro objetivo.",
            "Una pista aislada no acusa. Su valor surge al compararla con otras fuentes.",
        ),
        (
            "Oficina y pizarra\nHub físico donde se ordenan hallazgos, notas, conexiones y conclusiones.",
            "Volver a una memoria permite buscar algo que antes no podía interpretarse.",
            "Las conexiones son hipótesis editables y reciben respuesta narrativa según su coherencia.",
            "La acusación final evalúa una cadena completa, no la cantidad total de coleccionables.",
        ),
    ]
    for row_index, values in enumerate(rows, start=3):
        for column, value in enumerate(values):
            set_cell(table.rows[row_index].cells[column], value, size=7.8)

    values = [
        "VALORES NUCLEARES",
        "Contexto antes que acumulación",
        "Duda razonable antes que certeza automática",
        "Agencia y responsabilidad al concluir",
    ]
    for column, value in enumerate(values):
        set_cell(table.rows[6].cells[column], value, bold=True, size=8.0)


def fill_nouns(document: Document):
    table = document.tables[3]
    content = [
        ("Jugador / detective", "Posición, escena, herramienta activa, estado de interacción.", "Moverse, mirar, interactuar, inspeccionar, registrar, conectar y acusar."),
        ("Evidencia", "ID persistente, nombre, descripción, categoría, origen, foto y estado de descubrimiento.", "Ser observada, fotografiada, revelada, consultada, anotada y conectada."),
        ("Herramienta", "Tipo: mano, cámara, UV u objeto recogido; pose y estado activo.", "Equipar, guardar, elevar, encender, apuntar y utilizar."),
        ("Memoria", "Escena, requisitos de acceso, evidencias disponibles y cambios por conocimiento.", "Cargar, explorar, revisitar, actualizar y abandonar."),
        ("Teléfono", "Bloqueo, PIN opcional, apps, chats, llamadas y evidencias digitales.", "Recoger, desbloquear, navegar, leer y registrar información."),
        ("Expediente", "Víctima, implicados, información preliminar y actualizaciones del caso.", "Abrir, leer, cambiar página y actualizarse con el progreso."),
        ("Pizarra", "Posiciones de fotos/notas, conexiones e hipótesis persistentes.", "Colocar, arrastrar, anotar, conectar, desconectar y evaluar."),
        ("Conclusión", "Requisitos de método, acceso, motivo, tiempo y responsable.", "Desbloquear, presentar, objetar y validar."),
    ]
    set_cell(table.rows[0].cells[0], "Sustantivo", bold=True)
    set_cell(table.rows[0].cells[1], "Atributos relevantes para las mecánicas", bold=True)
    set_cell(table.rows[0].cells[2], "Verbos", bold=True)
    while len(table.rows) < len(content) + 1:
        table.add_row()
    for index, values in enumerate(content, start=1):
        for column, value in enumerate(values):
            set_cell(table.rows[index].cells[column], value, bold=column == 0, size=8.0)


def fill_mechanics(document: Document):
    rules = [
        "Si la mano apunta a un objeto interactuable dentro del alcance y no hay una UI modal abierta, E ejecuta su interacción contextual.",
        "Si se inspecciona un objeto, este se centra en el punto de inspección; el mouse rota el pivote y E restaura su transformación original.",
        "Si la cámara encuadra una EvidenceTarget visible y dentro de la distancia efectiva, el disparo registra una fotografía persistente.",
        "Si el haz UV alcanza una superficie revelable, solo los texeles expuestos acumulan revelado; al apagar la luz se conserva el progreso.",
        "Si se recoge el teléfono, se agrega al subinventario. Al equiparlo bloquea interacciones exteriores hasta guardarlo.",
        "Si el jugador abre mensajes o llamadas relevantes, cada hallazgo se registra como evidencia digital independiente.",
        "Si se conectan evidencias compatibles, se habilita una interpretación; si no son suficientes, la conexión queda como hipótesis y recibe una objeción.",
        "Si una acusación contiene responsable, método, acceso, motivo y tiempo respaldados, el caso puede cerrarse; de lo contrario continúa abierto.",
    ]
    table = document.tables[4]
    while len(table.rows) < len(rules):
        table.add_row()
    for index, rule in enumerate(rules):
        set_cell(table.rows[index].cells[0], rule, size=8.5)


def fill_ux(document: Document):
    table = document.tables[5]
    answers = {
        1: "Autosave por evidencia, conexiones y escena; recuperación de posición solo en memorias.",
        2: "Colliders, confirmación para borrar/salir y bloqueo de input detrás de interfaces modales.",
        4: "HUD contextual sin panel permanente; libreta, teléfono, expediente y pausa se excluyen entre sí.",
        5: "Texto reservado para narrativa, evidencia, objetivos y acciones que no puedan expresarse con iconos.",
        6: "TextMesh Pro, contraste alto, safe areas y layouts probados en 16:9 y resoluciones bajas.",
        7: "Rueda para acceso rápido, navegación por mouse/teclado y retorno directo con controles consistentes.",
        8: "Selector de memorias desde el monitor y herramientas de desarrollo para cambiar de escena durante pruebas.",
        9: "Carbón y papel para investigación; verde frío para sistema; rojo para relaciones; ámbar para objetivos.",
        10: "Canales separados para volumen maestro, efectos, voz/diálogo y ambiente.",
        12: "Animaciones de equipar/guardar, sonidos, retícula, flash, subtítulos y respuesta contextual.",
        13: "Movimiento tranquilo, sensibilidad configurable, FOV estable y objetos centrados al inspeccionar.",
        15: "Waypoints puntuales, retícula contextual, cambio de cursor y resaltado de objetivos actuales.",
        16: "Iconos de herramienta, apps de teléfono y fotografías conservan forma y función en todas las pantallas.",
        17: "Comandos comparten presentación; navegar, confirmar, volver e interactuar no intercambian significado.",
        19: "Introducción inicial define rol, víctima, expediente, objetivo de escena y condición de retorno.",
        20: "Cada sistema se aprende mediante una acción real y feedback inmediato, no mediante párrafos repetidos.",
        21: "Oficina: expediente y monitor. Casa: mano y cámara. Después: UV, teléfono, pizarra y acusación.",
        22: "Las ayudas no se reproducen simultáneamente y se guardan al completarse.",
        23: "Rol → interacción → fotografía → UV/teléfono → conexiones → acusación.",
        24: "La complejidad aumenta por contradicciones y relaciones, no por exigir más precisión motriz.",
        25: "Controles contextuales a la izquierda y pistas por inactividad sin repetir tutoriales completos.",
        27: "Curiosidad, reinterpretación de pistas, desbloqueo de memorias y construcción de una explicación propia.",
        28: "Feedback por descubrimiento y nueva relación; acusar sin sustento produce objeción, no pérdida total.",
        30: "Pruebas con jugadores que desconozcan el proyecto, registro de bloqueos y corrección por prioridad.",
        31: "Validar que el jugador comprenda que es detective, pueda explicar su hipótesis y diferencie evidencia de expediente.",
    }
    section_rows = {0, 3, 11, 14, 18, 26, 29}
    for index, row in enumerate(table.rows):
        if index in section_rows:
            set_cell(row.cells[0], normalize(row.cells[0].text), bold=True, size=8.8)
            continue
        set_cell(row.cells[1], answers.get(index, "Verificar en prueba de usuario."), size=7.8)


def fill_code_table(document: Document):
    table = document.tables[6]
    answers = [
        "Cumplido: se eliminaron separaciones excesivas y espacios finales.",
        "Cumplido: indentación uniforme de cuatro espacios.",
        "Cumplido: comentarios en responsabilidades y bloques cuya intención no es evidente.",
        "Cumplido: PascalCase para tipos/métodos y camelCase para campos privados serializados.",
        "En progreso: helpers compartidos para UI; próximos refactors dividirán los controladores más extensos.",
        "Cumplido por sistemas: evidencia, narrativa, pizarra, teléfono, guardado, UI y escena.",
        "Cumplido: GameManager persistente centraliza ciclo de vida y cambios de escena.",
        "Cumplido: Docs/CodeArchitecture.md contiene el diagrama Mermaid y reglas para extender el proyecto.",
    ]
    for index, answer in enumerate(answers):
        set_cell(table.rows[index].cells[1], answer, size=8.0)


def insert_concept_image(document: Document):
    _, end = remove_section_paragraphs(document, "2.3 Imagen", "2.4 Gesto")
    insert_paragraph_before(
        end,
        "La imagen resume el flujo visual del juego: memoria doméstica → oficina de análisis → evidencias contextualizadas → conclusión todavía incompleta. Los píxeles desplazados representan que las escenas son reconstrucciones y no reproducciones perfectas.",
    )
    paragraph = insert_paragraph_before(end, "")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(CONCEPT_IMAGE), width=Inches(6.65))
    caption = insert_paragraph_before(
        end,
        "Boceto abstracto de Archive: NULL. Casa investigada, oficina-hub, evidencias y cadena de deducción.",
        size=8.5,
    )
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sketch_title = insert_paragraph_before(
        end,
        "Bocetos rápidos de las diez situaciones interesantes",
        bold=True,
        size=9.5,
    )
    sketch_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sketch_paragraph = insert_paragraph_before(end, "")
    sketch_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sketch_paragraph.add_run().add_picture(str(SITUATION_SKETCHES), width=Inches(6.65))
    sketch_caption = insert_paragraph_before(
        end,
        "1. UV localizado · 2. Apps contradictorias · 3. Puerta e hilo · 4. Memoria revisitada · 5. Conexión objetada · 6. Documentos superpuestos · 7. Estado temporal · 8. Encuadre válido · 9. Evidencia contradictoria · 10. Acusación incompleta.",
        size=7.8,
    )
    sketch_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER


def build_document():
    document = Document(TEMPLATE)
    fill_cover(document)

    replace_section(
        document,
        "2.1 Deseo",
        "2.2 Situaciones",
        [
            "Archive: NULL debe expresar que investigar no es acumular objetos, sino reconstruir una verdad posible a partir de rastros incompletos, contradicciones y decisiones interpretativas.",
            "La experiencia buscada es la fantasía de ser un detective de reconstrucción forense: entrar en memorias digitales, observar una escena doméstica, registrar evidencia con herramientas concretas y volver a la oficina para construir una hipótesis defendible.",
            "El jugador debe llevarse la satisfacción de descubrir por qué una explicación aparentemente obvia no alcanza y de formular una acusación que pueda justificar con método, acceso, motivo y tiempo.",
            "El objetivo académico es integrar narrativa, interacción, persistencia, UI, sonido y programación en un caso jugable completo. La motivación del proyecto es demostrar que una investigación puede ser interesante por cómo obliga a pensar, no solo por ocultar una respuesta.",
        ],
    )
    remove_section_paragraphs(document, "2.2 Situaciones", "2.3 Imagen")
    fill_situations(document)
    insert_concept_image(document)
    replace_section(
        document,
        "2.4 Gesto",
        "2.5 Fantasía",
        [
            "El gesto predominante es observar con intención: desplazarse con WASD, orientar la vista con mouse y usar E sobre el elemento enfocado. La mano inspecciona; la cámara se eleva con F, encuadra y fotografía; la UV se enciende con F y debe barrerse físicamente sobre una superficie.",
            "G abre la rueda de herramientas. La libreta y la galería son pestañas separadas. En teléfono, expediente, pizarra y monitor se mantienen los mismos gestos de navegar, confirmar y volver mediante mouse o teclado.",
        ],
    )
    remove_section_paragraphs(document, "2.5 Fantasía", "2.6 Datos")
    fill_framework(document)
    replace_section(
        document,
        "2.6 Datos",
        "2.7 Key Features",
        [
            ("bullet", "Género: investigación narrativa y puzzle en primera persona."),
            ("bullet", "Temática: crimen doméstico, memoria reconstruida, evidencia plantada y contradicción."),
            ("bullet", "Plataforma y motor: PC, Unity 6, teclado y mouse."),
            ("bullet", "Modalidad: single player, offline."),
            ("bullet", "Duración objetivo del caso completo: 90 a 120 minutos."),
            ("bullet", "Público: jugadores de misterio que disfrutan observar, relacionar y justificar conclusiones."),
        ],
    )
    replace_section(
        document,
        "2.7 Key Features",
        "2.8 Descripción",
        [
            ("bold", "Memorias recontextualizables."),
            "Una escena puede revisitarse después de obtener conocimiento nuevo; objetos antes ambiguos adquieren interacciones y significado.",
            ("bold", "Herramientas forenses físicas y legibles."),
            "Mano para inspección, cámara con validación de encuadre, UV de revelado localizado y subinventario de dispositivos recogidos.",
            ("bold", "Teléfono como escena de investigación."),
            "PIN configurable, aplicaciones, chats, llamadas y hallazgos digitales que se convierten en evidencia persistente.",
            ("bold", "Pizarra de hipótesis libre con respuesta."),
            "Fotos y notas pueden moverse y conectarse; el juego diferencia una asociación posible de una relación demostrada.",
            ("bold", "Cierre por argumento."),
            "El caso no termina por encontrar un coleccionable ni por elegir un nombre: exige respaldar una cadena causal completa.",
        ],
    )
    replace_section(
        document,
        "2.8 Descripción",
        "2.9 Monetización",
        [
            "Archive: NULL es un juego de investigación narrativa en primera persona. El jugador encarna al Operador 253, detective de una división que reconstruye escenas mediante memorias digitales. El primer caso, La llave por dentro, investiga la muerte de Julián Herrera en una casa cerrada desde adentro. La disposición sugiere suicidio, pero el teléfono, los rastros químicos, documentos alterados y accesos secundarios contradicen esa lectura.",
            "El bucle alterna oficina y memoria. En la oficina se consulta el expediente, se selecciona una reconstrucción y se organizan fotos, notas y conexiones. En las memorias se inspeccionan objetos, se toman fotografías, se revelan rastros y se consultan dispositivos. La información obtenida actualiza el expediente y puede cambiar qué resulta interpretable al revisitar una escena.",
            "Cuando existe evidencia suficiente, el jugador presenta una acusación formada por responsable, método, acceso, motivo y secuencia temporal. Una respuesta correcta sin sustento no resuelve el caso.",
        ],
    )
    replace_section(
        document,
        "2.9 Monetización",
        "Espacio de juego",
        [
            "Pago único por una experiencia completa, sin microtransacciones. Una versión comercial podría ampliarse mediante casos independientes que reutilicen la oficina y los sistemas de investigación, pero el proyecto universitario prioriza cerrar un caso completo y evaluable.",
        ],
    )

    remove_section_paragraphs(document, "3.1 Sustantivos", "3.2 Mecánicas")
    fill_nouns(document)
    remove_section_paragraphs(document, "3.2 Mecánicas", "3.3 Requerimientos")
    fill_mechanics(document)
    replace_section(
        document,
        "3.3 Requerimientos",
        "3.4 Información",
        [
            ("bullet", "Toda evidencia debe poseer ID estable, nombre, descripción, procedencia y texto narrativo localizado."),
            ("bullet", "Las herramientas deben responder con animación, audio, retícula y controles contextuales."),
            ("bullet", "Una UI modal debe bloquear movimiento e interacción exterior y restaurarlos al cerrarse."),
            ("bullet", "El progreso debe persistir entre escenas: evidencias, fotos, notas, conexiones, conclusiones y posición en memorias."),
            ("bullet", "El tutorial debe avanzar por fases y no repetirse una vez completado."),
            ("bullet", "La experiencia debe funcionar en presets Bajo, Medio, Alto, Épico y Personalizado."),
        ],
    )
    replace_section(
        document,
        "3.4 Información",
        "Diseño de niveles",
        [
            "Bucle principal: EXPEDIENTE → SELECCIONAR MEMORIA → OBSERVAR/INSPECCIONAR → REGISTRAR EVIDENCIA → VOLVER A OFICINA → CONECTAR/ANOTAR → REVISITAR O DESBLOQUEAR MEMORIA → ACUSAR.",
            "La libreta guarda interpretación libre del jugador. La galería conserva fotografías y descripciones predeterminadas. El expediente aporta contexto institucional y se actualiza; la pizarra convierte hallazgos en hipótesis. Ninguno de estos sistemas reemplaza a los otros.",
        ],
    )

    replace_section(
        document,
        "4.1 Niveles",
        "4.1.1 Nombre",
        [
            "La progresión utiliza una oficina-hub y memorias investigables. El hub concentra contexto, selección, organización y cierre; cada memoria introduce una pregunta nueva y devuelve evidencia que modifica la lectura global del caso.",
        ],
    )
    replace_section(
        document,
        "4.1.1 Nombre",
        "4.1.2 Propósito",
        [
            ("bullet", "Oficina / Hub de investigación."),
            ("bullet", "Memoria 01: Casa Herrera — sala de estar y cocina."),
            ("bullet", "Memoria 02: Oficina de Víctor Salas — contratista y vecino."),
            ("bullet", "Revisita: Casa Herrera recontextualizada."),
            ("bullet", "Panel final de acusación."),
        ],
    )
    replace_section(
        document,
        "4.1.2 Propósito",
        "4.1.3 Descripción",
        [
            "La oficina presenta el rol de detective y evita que la primera memoria sea un espacio sin contexto. La Casa Herrera enseña observación, fotografía, UV y teléfono dentro de un entorno cotidiano. La oficina de Víctor aumenta la complejidad mediante comparación documental y evidencia de acceso. La revisita demuestra que investigar cambia el significado de lo ya visto. El panel final comprueba si el jugador puede argumentar, no si memorizó una respuesta.",
            "La emoción progresa de curiosidad y extrañeza a sospecha, contradicción y responsabilidad al acusar.",
        ],
    )
    replace_section(
        document,
        "4.1.3 Descripción",
        "4.1.4 Representación",
        [
            "Oficina: recorrido compacto entre puerta, expediente, monitor y pizarra. Casa: sala conectada con cocina, con líneas de visión que permiten reconocer objetos importantes sin señalarlos todos. Oficina de Víctor: espacio más denso y documental, organizado alrededor de escritorio, archivos y mesa de comparación.",
            "Los elementos narrativos principales se ubican en rutas naturales. Las evidencias opcionales amplían personajes o cronología, mientras las obligatorias sostienen método, acceso, motivo y tiempo.",
        ],
    )
    replace_section(
        document,
        "4.1.4 Representación",
        "4.2 Arte",
        [
            "Representación abstracta: rectángulo carbón = oficina; contorno verde = memoria disponible; círculo blanco = evidencia; línea roja continua = relación demostrada; línea roja punteada = hipótesis; cuadrados desplazados = reconstrucción incompleta; nodo ámbar = objetivo actual.",
        ],
    )
    replace_section(
        document,
        "4.2 Arte",
        "Experiencia de usuario",
        [
            "Dirección de realismo sobrio y doméstico. La oficina combina archivo analógico, monitor CRT y luz funcional; las memorias usan materiales cotidianos y pequeños defectos digitales. La iluminación debe guiar lectura espacial sin convertir toda evidencia en un objeto brillante.",
            "Paleta funcional: carbón para estructura, papel cálido para expediente, verde frío para sistemas, rojo apagado para conexiones y ámbar para objetivos. Los presets bajos deben conservar contraste, legibilidad y señales esenciales aunque reduzcan sombras, resolución o postprocesado.",
        ],
    )

    remove_section_paragraphs(document, "Experiencia de usuario", "Consideraciones sobre el código")
    fill_ux(document)
    remove_section_paragraphs(document, "Consideraciones sobre el código", "Extra")
    fill_code_table(document)
    replace_section(
        document,
        "Extra",
        "Pitch deck",
        [
            ("bullet", "Guardado automático persistente de evidencias, fotografías, notas, conexiones y estado de memoria."),
            ("bullet", "Revelado UV parcial mediante exposición localizada, no activación binaria del objeto."),
            ("bullet", "Teléfono diegético con navegación, PIN, mensajes y llamadas convertibles en evidencia."),
            ("bullet", "Sistema de gráficos configurable para equipos con GPU integrada."),
            ("bullet", "Localización español/inglés y controles reasignables."),
            ("bullet", "Auditor de assets de build y diagrama Mermaid de arquitectura."),
        ],
    )
    replace_section(
        document,
        "Pitch deck",
        "Para el examen final.",
        [
            "Pitch: una escena cerrada parece contar una historia completa, pero cada herramienta revela que alguien construyó esa historia para ser creída. Archive: NULL convierte al jugador en detective al exigirle observar, regresar, contradecir y defender una acusación propia.",
            "Prueba a desconocidos: verificar sin explicación externa si comprenden el rol, encuentran el expediente, acceden a la primera memoria, usan mano/cámara/UV/teléfono, regresan a la oficina, conectan evidencia y pueden explicar qué creen que ocurrió. Registrar tiempo, bloqueos, interpretaciones y controles olvidados.",
        ],
    )

    document.core_properties.title = "Archive: NULL - Game Design Document V4"
    document.core_properties.subject = "Proyecto de Integración"
    document.save(OUTPUT)

    # Guard against artifacts produced by the previous Markdown conversion script.
    reopened = Document(OUTPUT)
    all_text = "\n".join(p.text for p in reopened.paragraphs)
    all_text += "\n" + "\n".join(cell.text for table in reopened.tables for row in table.rows for cell in row.cells)
    forbidden = ["param($m)", "ACTUALIZACIÓN V4", "Situación 11:", "4.1 La evidencia"]
    found = [value for value in forbidden if value in all_text]
    if found:
        raise RuntimeError(f"Forbidden legacy content found: {found}")

    print(f"Created {OUTPUT}")


if __name__ == "__main__":
    build_document()
